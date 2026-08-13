"""Build end-user taxonomy exports (Excel / Word) from path rows."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Font
from docx import Document


def _rows_from_path_items(items: list[Any]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for item in items:
        rows.append(
            {
                "Zone": item.zone,
                "Department": item.department,
                "Category": item.category,
                "Subcategory": item.subcategory,
                "Active": "Yes" if item.is_active else "No",
                "Full path": item.full_path,
            }
        )
    return rows


def build_excel(items: list[Any]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Taxonomy"
    headers = ["Zone", "Department", "Category", "Subcategory", "Active", "Full path"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    for row in _rows_from_path_items(items):
        ws.append([row[h] for h in headers])
    for col in ws.columns:
        width = max(len(str(c.value or "")) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max(width + 2, 12), 48)
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_word(items: list[Any]) -> bytes:
    doc = Document()
    doc.add_heading("Retail Taxonomy Export", level=1)
    doc.add_paragraph(
        "Hierarchy paths (Zone → Department → Category → Subcategory)."
    )
    rows = _rows_from_path_items(items)
    headers = ["Zone", "Department", "Category", "Subcategory", "Active", "Full path"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, h in enumerate(headers):
            table.rows[r_idx].cells[c_idx].text = row[h]
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
